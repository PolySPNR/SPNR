import json
from docx import Document
import sys


arguments = sys.argv[1:]
max_year = 0
sort_tables = [['монография', 0, 0], ['научный труд', 0, 0], ['учебник', 0, 0], ['учебное пособие', 0, 0],
               ['методическое указание', 0, 0], ['статья', 0, 0], ['конференция', 0], ['патент', 0],
               ['программный продукт', 0], ['лицензия', 0], ['выставка', 0]]
journals_publish_name = []
journals_authors_names = []
journals_output_data = []
journals_count_pages = []
with open(str(arguments[0]), 'r', encoding='utf-8') as fh:
    data = json.load(fh)
author_id_number = 0
for info in data:
    publication_type = info["PublicationInfo"]["Тип"]
    count_pages = info["PublicationInfo"]["Страницы"]
    if int(info["PublicationInfo"]["Год"]) > max_year:
        max_year = int(info["PublicationInfo"]["Год"])
    for type_table in sort_tables[:6]:
        count_pages_int = 0
        if type_table[0] in publication_type and type_table[0] == 'статья':
            count_pages_int = (int(count_pages[3]) * 10 + int(count_pages[4]) - int(count_pages[0]) * 10 - int(
                count_pages[1])) / 16
            type_table[2] += count_pages_int
            type_table[1] += 1
            journals_publish_name.append(info["WorkName"])
            author_id_number += 1
            for author in info["Authors"]:
                journals_authors_names.append([author["Name"], author_id_number])
            journals_output_data.append(info["PublicationMeta"]["ЖУРНАЛ"])
            journals_count_pages.append(count_pages_int)


document = Document()
paragraph_1 = document.add_paragraph()
paragraph_1.add_run('Приложение 2')
paragraph_1.alignment = 2
paragraph_2 = document.add_paragraph()
paragraph_2.add_run('к Распоряжению № _____ от ______ г.')
paragraph_2.alignment = 2
document.add_heading('РЕЗУЛЬТАТИВНОСТЬ НАУЧНО-ИССЛЕДОВАТЕЛЬСКОЙ ДЕЯТЕЛЬНОСТИ КАФЕДРЫ «Информационная безопасность» В' +
                     ' ' + str(max_year) + ' ' + 'ГОДУ', 1)
all_publish = document.add_table(rows=len(sort_tables[:6]) + 1, cols=2)
all_publish.style = 'Table Grid'
heading_all_publish = all_publish.rows[0].cells
heading_all_publish[0].paragraphs[0].add_run('Показатель').bold = True
heading_all_publish[1].paragraphs[0].add_run('Количество').bold = True
iterator = 1
for info in sort_tables[:6]:
    all_publish_row_data = all_publish.rows[iterator].cells
    all_publish_row_data[0].paragraphs[0].add_run(info[0]).bold = True
    all_publish_row_data[1].paragraphs[0].add_run(str(info[1]) + ' ' + '(' + str(info[2]) + ' ' + 'п.л.'
                                                  + ')').bold = False
    iterator += 1
paragraph_3 = document.add_paragraph()
paragraph_3.add_run('Зав. кафедрой')
paragraph_4 = document.add_paragraph()
paragraph_4.add_run('«Информационная безопасность»    _____________________________Н.В. Федоров')
document.add_page_break()
paragraph_5 = document.add_paragraph()
paragraph_5.add_run('Расшифровка к приложению 2').bold = True
paragraph_5.alignment = 2
document.add_heading('РЕЗУЛЬТАТИВНОСТЬ НИР КАФЕДРЫ', 1)
table_pointer = 0
for info in sort_tables:
    table_pointer += 1
    table_paragraph = document.add_paragraph(info[0] + ' ' + '(Таблица 2.' + str(table_pointer) + ')')
    table = document.add_table(rows=info[1] + 2, cols=5)
    table.style = 'Table Grid'
    sum_count_list = 0
    heading = table.rows[0].cells
    heading[0].paragraphs[0].add_run('№№').bold = True
    heading[1].paragraphs[0].add_run('Название статьи').bold = True
    heading[2].paragraphs[0].add_run('ФИО авторов и соавторов с указанием кафедр (организаций)').bold = True
    heading[3].paragraphs[0].add_run('Выходные данные').bold = True
    heading[4].paragraphs[0].add_run('Кол-во печатных листов').bold = True
    if info[1] == 0:
        pass
    else:
        iterator = 0
        data_pointer = 1
        for i in range(1, info[1] + 1):
            row_data = table.rows[i].cells
            for j in range(5):
                if heading[j].text == '№№':
                    row_data[j].paragraphs[0].add_run(str(data_pointer) + '.').bold = True
                elif heading[j].text == 'Название статьи':
                    row_data[j].paragraphs[0].add_run(str(journals_publish_name[iterator]))
                elif heading[j].text == 'ФИО авторов и соавторов с указанием кафедр (организаций)':
                    for author in journals_authors_names:
                        if author[1] == data_pointer:
                            row_data[j].paragraphs[0].add_run(str(author[0]) + ' ')
                elif heading[j].text == 'Выходные данные':
                    for output_data in journals_output_data[iterator]:
                        row_data[j].paragraphs[0].add_run(str(output_data + ' '))
                elif heading[j].text == 'Кол-во печатных листов':
                    sum_count_list += float(journals_count_pages[iterator])
                    row_data[j].paragraphs[0].add_run(str(journals_count_pages[iterator]))
            data_pointer += 1
            iterator += 1
        heading = table.rows[len(data) + 1].cells
        heading[0].paragraphs[0].add_run(str(info[1] + 1) + '.').bold = True
        heading[3].paragraphs[0].add_run('Итого (п.л.):').bold = True
        heading[4].paragraphs[0].add_run(str(sum_count_list)).bold = True
document.save('таблица.docx')
